# ascii generator
import argparse
from PIL import Image
from docx import Document

# 10 levels of gray
# the best: 4, 11, 6, 13, 14, 2, 10
gscale0 = "$@B%8&WM#*oahkbdpqwmZO0QLCJUYXzcvunxrjft/\\|()1{}[]?-_+~<>i!lI;:,\"^`'. "

gscale1 = "$@B%8&WM*oakbdpqwmZO0QLCJUYXzcvunxrjft/\\|()1{}[]?-_+~><i!lI;:,\"^`'. "

gscale2 = "MW&8#%xo*=+~:-. "

gscale3 = "%&#oakbdqpmwZO0QJUYXzcvnxvt/|)(][}{?1Iil!*+=~-;:,'`. "

gscale4 = "MW&@$%#*+=-:. "

gscale5 = "MWoakbdqpmwZOQLCJUYXzcvnuxrjft!li;:,.'` "

gscale6 = "@%#*+=-:. "

gscale7 = "@#%WM&$*=+~^-/;:,. "

gscale8 = "@#%WM&$B*=+/~^:;,. "

gscale9 = "@#%WM&$=*+~-^:/;,.'` "

gscale10 = "@#%WM&$*+=~^:-., "

gscale11 = "@#%WM&$B8*oa+=~^:;-,. "

gscale12 = "@#MW&%$B8o*a+=~^:;-,. "

gscale13 = "@#&*=+-:. "

gscale14 = "█▓▒░:. "

gscale15 = "@WM$*o+=-:. "

gscale16 = "█@▓W&%MØ$ΦΔ8o╬╣xo=/~:;,'`. "

gscale17 = "@#WɅÆ&%╬$BM8oah*=+~:;,.` "

gscale18 = "█▒░▓█▒░▓▓▒░██░▒▒/\\/\\/\\\\/=~^^-:. "

gscale19 = "ΩΣ@#&%∞≡≠≃≈∴πλ+=~–:,'`. "

gscale20 = "█◉⨀✷✧⟁⟁⨂⨁⊚⚛*⋆∘·˚˖,˙. "

gscale21 = "█▓▒░╬╩╦═╚╝╔╗╠╣═╬++--::.. "

gscale22 = "█⯈⯊⨀◉✷☢⟁⨂⊚☯⚛⦿∘⋆·˚˖,˙. "

gscale23 = "█⨸☠⩚⯈⚿⟁⨃⩰⚛⩪⋆˚᯽᯾᯿´'˙. "

gscale24 = "⬛█⯈⨀◉✶✧✷⟁⨂⊚⚛⦿∘⋆·˚˖,˙.⬜"

gscale25 = "█⯈⯊⨀⨸☠⚿◉✷✶☢☯⟁⨃⨂⊚⚛⦿⩚⩰⩪∘⋆·˚᯽᯾᯿´'˙.,.⠀ "

gscale26 = "█⨀⨸Φ⟁Ϟ⧫⩚⚿ℵ⩰⚛∞⦿∴∘⋆·˚˖,˙.⠀"

gscale27 = "█☠☢⛧⨸⚿⊗⨂⚛⟁᯾∴⋆⋇⌖·˖,˙.⠀"

gscale28 = "⣿⡿⣾⣽⣻⢿⢾⢽⣷⣯⣟⣠⡄⡀⠄⠂⠁⠀"

gscale29 = "█⟁⧖⧗⧚☍⚛⟊⟐⩰∷⋆·⠂⠁⠀"

gscale30 = "█⨀𓂀⚱⚛⚿⟁⧫⦿⌖⋆·˚˖,˙.⠀"

gscale31 = "█▓⣿⯈⨸☠⚿⛧⚛⛩⩚⧫∑♒♾☯⚰⛓✠⌖∷∘⋆˚᯾᯿,˙.⠀"

greyscales = [
    gscale0,
    gscale1,
    gscale2,
    gscale3,
    gscale4,
    gscale5,
    gscale6,
    gscale7,
    gscale8,
    gscale9,
    gscale10,
    gscale11,
    gscale12,
    gscale13,
    gscale14,
    gscale15,
    gscale16,
    gscale17,
    gscale18,
    gscale19,
    gscale20,
    gscale21,
    gscale22,
    gscale23,
    gscale24,
    gscale25,
    gscale26,
    gscale27,
    gscale28,
    gscale29,
    gscale30,
    gscale31,  # Singularity Core
]

from docx import Document
from docx.shared import Pt, Inches, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_LINE_SPACING


def save_ascii_art_to_docx(ascii_lines, filename="ascii_art.docx"):
    doc = Document()

    # Define A4 size in EMUs
    A4_WIDTH = Emu(8.27 * 914400)  # ~7,566,288 EMU
    A4_HEIGHT = Emu(11.69 * 914400)  # ~10,473,779 EMU

    for section in doc.sections:
        # Set page size to A4
        section.page_width = A4_WIDTH
        section.page_height = A4_HEIGHT

        # Set all margins to zero
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)

    for line in ascii_lines:
        para = doc.add_paragraph()
        run = para.add_run(line)

        # Font settings
        font = run.font
        font.name = "Courier New"
        font.size = Pt(5)

        # Force eastAsia font override safely
        rpr = run._element.rPr
        if rpr is not None:
            rFonts = rpr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rpr.append(rFonts)
            rFonts.set(qn("w:eastAsia"), "Courier New")

        # Minimal spacing between lines
        para_format = para.paragraph_format
        para_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        para_format.line_spacing = Pt(5)
        para_format.space_before = Pt(0)
        para_format.space_after = Pt(0)

    doc.save(filename)


def average(hist):
    suma = 0
    division = 0
    for i, amount in enumerate(hist):
        suma += i * amount
        division += amount

    if division == 0:
        return 0
    return suma / division


def generate_new_image(img, cols, gscale, x):
    image = Image.open(img).convert("L")
    width = image.size[0]
    height = image.size[1]

    if width < cols and x < 1:
        raise ValueError("to small picture")

    if x != 0:
        scale = x
        cols = int(width / scale)
    else:
        scale = int(width / cols)

    sub = 2

    rows = int(height / scale / sub)
    ascii = []

    for y in range(rows):
        y1 = y * scale * sub
        y2 = y * scale * sub + scale

        output = ""
        for x in range(cols):
            x1 = x * scale
            x2 = x * scale + scale

            crop = image.crop((x1, y1, x2, y2))
            gr_scale = gscale[
                int((average(crop.histogram()) * (len(gscale) - 1)) / 255)
            ]
            output += gr_scale
        ascii.append(output)
    return ascii


def main():
    desc = "convert images to ascii art."
    parser = argparse.ArgumentParser(description=desc)

    # add arguments to parser
    parser.add_argument("--img", dest="img", help="source image", required=True)
    parser.add_argument(
        "--cols",
        dest="cols",
        help="your terminal width",
        type=int,
        default=120,
        required=False,
    )
    parser.add_argument(
        "--out",
        dest="out",
        help="your output file name or terminal if not mentioned",
        default="terminal",
        required=False,
    )
    parser.add_argument("--more", dest="more", help="grey scale", type=int, default=4)
    parser.add_argument(
        "--x",
        dest="x",
        help="your scale",
        type=int,
        default=0,
        required=False,
    )

    args = parser.parse_args()

    img = args.img
    cols = args.cols
    output = args.out
    more = args.more
    x = args.x

    if more < len(greyscales):
        gscale = greyscales[more]
    else:
        print("wrong")

    new_image = generate_new_image(img, cols, gscale, x)

    if output == "terminal":
        for line in new_image:
            print(line)
    elif ".docx" in output:
        save_ascii_art_to_docx(new_image, output)
    else:
        with open(f"{output}", "w") as text_file:
            text_file.writelines("\n".join(new_image))


if __name__ == "__main__":
    main()
